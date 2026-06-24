# -*- coding: utf-8 -*-
"""Sinh bao cao Word cho do an OOP - He thong quan ly thu vien."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r"E:\LTHDT_OOP\TL CK\BaoCao_QuanLyThuVien_OOP.docx"
ASSETS = r"E:\LTHDT_OOP\TL CK\report_assets"

doc = Document()

# ---------- Style co ban ----------
base = doc.styles["Normal"]
base.font.name = "Times New Roman"
base.font.size = Pt(13)
base._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
doc.styles["Normal"].paragraph_format.space_after = Pt(6)
doc.styles["Normal"].paragraph_format.line_spacing = 1.4


def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def para(text="", size=13, bold=False, italic=False, align=None, color=None,
         space_before=0, space_after=6, font="Times New Roman"):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    return p


def heading(text, level):
    sizes = {1: 16, 2: 14, 3: 13}
    colors = {1: (0x0B, 0x3D, 0x91), 2: (0x12, 0x55, 0xA0), 3: (0x1F, 0x4E, 0x79)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = RGBColor(*colors[level])
    return p


def code_block(code, caption=None):
    """Khung code nen xam, font monospace."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F5F7")
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), "C0C4CC")
        borders.append(e)
    tcPr.append(borders)
    cell.paragraphs[0].text = ""
    first = True
    for line in code.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(10)
    if caption:
        para(caption, size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(0x55, 0x55, 0x55), space_after=10)


def output_block(text, caption=None):
    """Khung ket qua console nen den."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "1E1E1E")
    cell.paragraphs[0].text = ""
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    if caption:
        para(caption, size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(0x55, 0x55, 0x55), space_after=10)


def add_image(name, caption=None, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(ASSETS, name), width=Inches(width))
    if caption:
        para(caption, size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(0x55, 0x55, 0x55), space_after=12)


def make_table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "0B3D91")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11.5)
    if widths:
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl


# ==================== TRANG BIA ====================
para("TRƯỜNG ĐẠI HỌC ...", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=2)
para("KHOA CÔNG NGHỆ THÔNG TIN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=40)

para("BÁO CÁO ĐỒ ÁN CUỐI KỲ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18,
     color=(0x0B, 0x3D, 0x91), space_after=6)
para("MÔN: LẬP TRÌNH HƯỚNG ĐỐI TƯỢNG", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=30)

para("ĐỀ TÀI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=4)
para("HỆ THỐNG QUẢN LÝ THƯ VIỆN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20,
     color=(0xB0, 0x30, 0x20), space_after=50)

info = doc.add_table(rows=4, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = [("Lớp / Nhóm", "Thứ 6 - Nhóm 13 (T6_N13)"),
          ("Giảng viên hướng dẫn", "..............................."),
          ("Sinh viên thực hiện", "..............................."),
          ("Năm học", "2025 - 2026")]
for i, (k, v) in enumerate(labels):
    c0 = info.cell(i, 0); c1 = info.cell(i, 1)
    r = c0.paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(13); r.font.name = "Times New Roman"
    r = c1.paragraphs[0].add_run(v); r.font.size = Pt(13); r.font.name = "Times New Roman"

doc.add_page_break()

# ==================== MUC LUC ====================
heading("MỤC LỤC", 1)
toc = [
    "Chương 1. Giới thiệu đề tài và sự phù hợp với mô hình OOP",
    "Chương 2. Phân tích và thiết kế hệ thống",
    "Chương 3. Cài đặt các khối chức năng (code + kết quả)",
    "Chương 4. Kết luận",
]
for t in toc:
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].bold = True
doc.add_page_break()

# ==================== CHUONG 1 ====================
heading("CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI VÀ SỰ PHÙ HỢP VỚI MÔ HÌNH OOP", 1)

heading("1.1. Giới thiệu đề tài", 2)
para("Đề tài của nhóm là xây dựng một Hệ thống quản lý thư viện viết bằng C++ theo mô hình "
     "lập trình hướng đối tượng (OOP). Hệ thống mô phỏng đầy đủ các nghiệp vụ thực tế của một "
     "thư viện: quản lý con người (độc giả, nhân viên, quản trị viên), quản lý tài liệu "
     "(giáo trình, sách tham khảo, sách khác, báo/tạp chí, bài nghiên cứu) và quản lý các "
     "giao dịch (mượn, trả, nhập, xuất/thanh lý).")
para("Toàn bộ logic nghiệp vụ được tổ chức thành các lớp (class) có quan hệ kế thừa, đóng gói "
     "dữ liệu và sử dụng đa hình để xử lý linh hoạt từng loại đối tượng. Ngoài phiên bản chạy "
     "trên console, nhóm còn phát triển thêm một bản demo web (thư mục web/) để minh họa trực quan.")

heading("1.2. Mục tiêu của đề tài", 2)
bullet("Mô hình hóa các thực thể trong thư viện thành các lớp OOP rõ ràng, dễ mở rộng.")
bullet("Áp dụng đầy đủ 4 đặc tính cốt lõi của OOP: trừu tượng, đóng gói, kế thừa, đa hình.")
bullet("Xử lý đúng các nghiệp vụ: kiểm tra điều kiện mượn, tính hạn mượn theo từng đối tượng, "
       "tính tiền phạt trễ hạn, cập nhật tồn kho khi nhập/xuất.")
bullet("Bảo đảm tính đúng đắn của dữ liệu đầu vào (kiểm tra ngày tháng, số lượng, mã trùng...).")

heading("1.3. Đề tài đáp ứng các tiêu chí của một hệ thống OOP chuẩn", 2)
para("Một hệ thống được xem là “chuẩn OOP” khi thể hiện được đầy đủ các đặc tính nền tảng "
     "của lập trình hướng đối tượng. Bảng dưới đây đối chiếu từng tiêu chí với cách đề tài đáp ứng:")

make_table(
    ["Tiêu chí OOP", "Yêu cầu", "Đề tài đáp ứng như thế nào"],
    [
        ["Trừu tượng\n(Abstraction)",
         "Có lớp trừu tượng / hàm ảo thuần túy làm khung thiết kế chung.",
         "3 lớp cha Nguoi, TaiLieu, GiaoDich đều là abstract class với hàm ảo thuần túy "
         "(= 0), không thể khởi tạo trực tiếp."],
        ["Đóng gói\n(Encapsulation)",
         "Dữ liệu được bảo vệ, truy xuất qua getter/setter.",
         "Thuộc tính để private/protected; mọi truy xuất đều thông qua phương thức công "
         "khai, kèm kiểm tra ràng buộc (vd capNhatSoLuong không cho âm)."],
        ["Kế thừa\n(Inheritance)",
         "Tái sử dụng và tổ chức mã theo phân cấp.",
         "Hệ thống dùng kế thừa nhiều cấp: Nguoi → NguoiDoc → SinhVien/VienChuc; "
         "TaiLieu → Sach → GiaoTrinh/SachThamKhao/SachKhac."],
        ["Đa hình\n(Polymorphism)",
         "Cùng một lời gọi cho ra hành vi khác nhau theo loại đối tượng.",
         "Hàm ảo tinhHanMuon(), hienThiThongTin(), thucHienGiaoDich() được override; "
         "cùng tinhHanMuon nhưng SV và VC cho hạn khác nhau."],
        ["Khuôn mẫu\n(Template)",
         "(Mở rộng) Tái sử dụng cấu trúc cho nhiều kiểu dữ liệu.",
         "Lớp template DanhSach<T> dùng chung cho cả DanhSach<Nguoi>, "
         "DanhSach<TaiLieu>, DanhSach<GiaoDich>."],
        ["Quản lý bộ nhớ",
         "An toàn, tránh rò rỉ.",
         "Sử dụng smart pointer (shared_ptr) và destructor ảo (virtual ~) để hủy "
         "đúng đối tượng lớp con."],
    ],
    widths=[1.3, 2.2, 3.0],
)
para("")
para("Như vậy đề tài không chỉ đáp ứng mà còn minh họa rõ nét tất cả các trụ cột của OOP, "
     "phù hợp để làm đồ án cuối kỳ cho môn học.", italic=True)

doc.add_page_break()

# ==================== CHUONG 2 ====================
heading("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
para("Chương này trình bày phần phân tích – thiết kế hướng đối tượng đã báo cáo với giảng viên: "
     "xác định các đối tượng, sơ đồ phân cấp lớp và cách áp dụng các đặc tính OOP.")

heading("2.1. Xác định các đối tượng", 2)
para("Hệ thống xoay quanh 3 nhóm đối tượng lớn:")
bullet("Con người: mỗi cá nhân tương tác với thư viện (độc giả, nhân viên, quản trị viên).", "Nhóm 1 – ")
bullet("Tài liệu: mỗi ấn phẩm lưu hành trong thư viện (sách, báo/tạp chí, bài nghiên cứu).", "Nhóm 2 – ")
bullet("Giao dịch: mỗi nghiệp vụ làm thay đổi trạng thái thư viện (mượn, trả, nhập, xuất).", "Nhóm 3 – ")

heading("2.2. Sơ đồ phân cấp lớp", 2)
hierarchy = """Nguoi (abstract)
 |-- NguoiDoc (abstract)
 |    |-- SinhVien        (gioi han 5 sach)
 |    +-- VienChuc        (gioi han 10 sach)
 |-- NhanVien
 +-- QuanTriVien

TaiLieu (abstract)
 |-- Sach
 |    |-- GiaoTrinh
 |    |-- SachThamKhao
 |    +-- SachKhac
 |-- BaoTapChi            (chi doc tai cho)
 +-- BaiNghienCuu         (chi doc tai cho)

GiaoDich (abstract)
 |-- Muon
 |-- Tra
 |-- Nhap
 +-- Xuat

DanhSach<T>  (lop khuon mau quan ly danh sach)"""
code_block(hierarchy, "Hình 2.1. Sơ đồ phân cấp kế thừa các lớp trong hệ thống")

heading("2.3. Bảng mô tả chi tiết các lớp đối tượng", 2)
make_table(
    ["Lớp", "Thuộc tính chính", "Phương thức tiêu biểu", "Đặc tính OOP"],
    [
        ["Nguoi (cha)", "maNguoi, hoTen, ngaySinh, gioiTinh, soDienThoai, diaChi",
         "hienThiThongTin() [ảo]", "Trừu tượng, Đóng gói"],
        ["NguoiDoc (cha)", "ngayDangKy, ngayHetHan, soSachDangMuon, gioiHanSachMuon",
         "duocPhepMuonSach(), tangSoSachMuon()", "Kế thừa, Trừu tượng"],
        ["SinhVien", "mssv", "hanMuonGiaoTrinh/ThamKhao/Khac()", "Kế thừa, Đa hình"],
        ["VienChuc", "mscb", "hanMuon... (hạn dài hơn SV)", "Kế thừa, Đa hình"],
        ["NhanVien", "chucVu, luong, caLamViec", "tinhLuong(), phanCaLam()", "Kế thừa, Đa hình"],
        ["QuanTriVien", "username, password, quyenTruyCap", "phanQuyen(), khoiPhucMatKhau()", "Kế thừa, Đa hình"],
        ["TaiLieu (cha)", "maTaiLieu, tenTaiLieu, namXuatBan, soLuong, trangThai, tacGia...",
         "capNhatSoLuong(), tinhHanMuon() [ảo]", "Trừu tượng, Đóng gói"],
        ["Sach", "(kế thừa TaiLieu)", "quanLySoLuong()", "Kế thừa"],
        ["GiaoTrinh", "maMonHoc, boMon", "kiemTraHocPhan(), tinhHanMuon()", "Kế thừa, Đa hình"],
        ["SachThamKhao", "(kế thừa Sach)", "tinhHanMuon()", "Kế thừa, Đa hình"],
        ["SachKhac", "loaiSachKhac", "tinhHanMuon()", "Kế thừa, Đa hình"],
        ["BaoTapChi", "soPhatHanh, thangPhatHanh", "docTaiCho(), coTheMuonVe()=false", "Kế thừa, Đa hình"],
        ["BaiNghienCuu", "coQuanChuQuan, linhVuc", "traCuuNghienCuu(), coTheMuonVe()=false", "Kế thừa, Đa hình"],
        ["GiaoDich (cha)", "maGiaoDich, ngayGiaoDich, maNhanVienThucHien",
         "thucHienGiaoDich() [ảo]", "Trừu tượng, Đóng gói"],
        ["Muon", "maNguoiDoc, maTaiLieu, ngayHenTra", "kiemTraQuaHan(), ghiNhanNgayMuon()", "Kế thừa, Đa hình"],
        ["Tra", "tienPhat, ngayHenTra", "tinhTienPhat(), xacNhanTraSach()", "Kế thừa, Đa hình"],
        ["Nhap", "nhaCungCap, soLuongNhap, donGiaNhap", "tinhTongChiPhi()", "Kế thừa, Đa hình"],
        ["Xuat", "lyDoXuat, soLuongXuat, donGiaThanhLy", "tinhDoanhThuThanhLy(), loaiBoTaiLieu()", "Kế thừa, Đa hình"],
    ],
    widths=[1.1, 2.0, 2.1, 1.3],
)

heading("2.4. Áp dụng các đặc tính OOP", 2)

heading("a) Tính trừu tượng (Abstraction)", 3)
para("Ba lớp cha Nguoi, TaiLieu, GiaoDich được khai báo là lớp trừu tượng: chúng chứa ít nhất "
     "một hàm ảo thuần túy (= 0) nên không thể tạo đối tượng trực tiếp, chỉ đóng vai trò “bản "
     "thiết kế” cho các lớp con.")

heading("b) Tính đóng gói (Encapsulation)", 3)
para("Các thuộc tính nhạy cảm (mã, tên, số lượng, mật khẩu...) đều để private/protected. Bên "
     "ngoài chỉ truy xuất qua getter/setter, qua đó có thể lồng thêm ràng buộc kiểm tra "
     "(ví dụ không cho số lượng xuống dưới 0).")

heading("c) Tính kế thừa (Inheritance)", 3)
para("Hệ thống dùng kế thừa phân cấp nhiều cấp. Lớp con tái sử dụng toàn bộ thuộc tính – phương "
     "thức của lớp cha và chỉ bổ sung phần riêng (SinhVien thêm mssv, VienChuc thêm mscb...).")

heading("d) Tính đa hình (Polymorphism)", 3)
para("Áp dụng mạnh nhất ở quy tắc tính hạn mượn: cùng lời gọi taiLieu->tinhHanMuon(nguoiDoc) "
     "nhưng kết quả khác nhau theo từng loại tài liệu và từng loại độc giả:")
make_table(
    ["Loại tài liệu", "Sinh viên", "Viên chức"],
    [
        ["Giáo trình", "180 ngày", "180 ngày"],
        ["Sách tham khảo", "15 ngày", "90 ngày"],
        ["Sách khác", "7 ngày", "30 ngày"],
        ["Báo/tạp chí, Bài nghiên cứu", "Chỉ đọc tại chỗ (không mượn về)", "Chỉ đọc tại chỗ"],
    ],
    widths=[3.0, 1.7, 1.7],
)

heading("e) Khuôn mẫu (Template) – mở rộng", 3)
para("Lớp DanhSach<T> được viết một lần và dùng lại cho mọi loại danh sách (Nguoi, NguoiDoc, "
     "TaiLieu, GiaoDich), tránh viết lặp các cấu trúc danh sách giống nhau.")

heading("2.5. Phân tích đối tượng và mô hình dữ liệu ứng dụng web", 2)
para("Ngoài bản C++ console, nhóm hiện thực hệ thống thành một ứng dụng web (app.js). Bản web "
     "giữ nguyên mô hình hướng đối tượng nhưng lưu dữ liệu dưới dạng một đối tượng state JSON "
     "duy nhất trong localStorage của trình duyệt (khóa oop-library-web-v5). Toàn bộ nghiệp vụ "
     "thao tác trực tiếp trên state này.")
para("Qua phân tích mã nguồn, hệ thống gồm 7 nhóm đối tượng (thực thể) chính sau:")
make_table(
    ["Đối tượng (state)", "Vai trò", "Thuộc tính tiêu biểu"],
    [
        ["accounts", "Tài khoản đăng nhập, gắn với 3 vai trò admin / staff / reader",
         "username, password, role, name, title, personId, readerId"],
        ["readers", "Độc giả: Sinh viên (giới hạn 5) hoặc Viên chức (giới hạn 10)",
         "id, name, type, birth, gender, phone, address, code (MSSV/MSCB), "
         "registered, expires, borrowed, limit"],
        ["staffs", "Nhân viên / thủ thư", "id, name, chức vụ, lương, ca làm, username"],
        ["admins", "Quản trị viên", "id, name, username, permission"],
        ["documents", "Tài liệu thuộc 5 loại (kind), mỗi loại có thêm trường extra riêng",
         "id, title, kind, year, quantity, author, publisher, category, extra{...}, "
         "coverImage/fileUrl"],
        ["loans", "Phiếu mượn đang mở (mã PM…)",
         "id, readerId, documentId, borrowDate, dueDate"],
        ["transactions", "Giao dịch (mã GD…): Mượn / Trả / Nhập / Xuất",
         "id, type, date, staffId, readerId, documentId, quantity, unitPrice, amount"],
    ],
    widths=[1.5, 2.5, 2.5],
)
para("")
add_image("fig_object_model.png",
          "Hình 2.2. Sơ đồ đối tượng và quan hệ của ứng dụng web (mô hình dữ liệu)", width=6.5)

para("Mỗi đối tượng dữ liệu của bản web tương ứng với một lớp trong thiết kế OOP C++:", bold=True)
make_table(
    ["Đối tượng web (app.js)", "Lớp OOP tương ứng (C++)"],
    [
        ["accounts (role=reader/staff/admin)", "QuanTriVien / phân quyền tài khoản"],
        ["readers (type SV/VC)", "NguoiDoc → SinhVien / VienChuc"],
        ["staffs", "NhanVien"],
        ["admins", "QuanTriVien"],
        ["documents (kind)", "TaiLieu → GiaoTrinh / SachThamKhao / SachKhac / BaoTapChi / BaiNghienCuu"],
        ["loans", "Phiếu mượn (PhieuMuonDangMo)"],
        ["transactions (type)", "GiaoDich → Muon / Tra / Nhap / Xuat"],
    ],
    widths=[3.2, 3.3],
)

heading("2.6. Sơ đồ luồng dữ liệu (Data Flow Diagram)", 2)
para("Sơ đồ luồng dữ liệu mô tả dữ liệu di chuyển giữa các tác nhân, tiến trình và kho dữ liệu. "
     "Hệ thống có 3 tác nhân ngoài (độc giả, nhân viên, quản trị viên) và một kho dữ liệu là "
     "localStorage của trình duyệt.")
add_image("fig_dfd_context.png",
          "Hình 2.3. DFD mức ngữ cảnh – tương tác giữa 3 tác nhân và hệ thống", width=6.3)
para("Phân rã mức 1 chia hệ thống thành 5 tiến trình chính, ghi/đọc trên 5 kho dữ liệu (D1–D5) "
     "đều nằm trong state localStorage:")
add_image("fig_dfd_level1.png",
          "Hình 2.4. DFD mức 1 – các tiến trình và kho dữ liệu", width=6.6)

heading("2.7. Tổng hợp chức năng hệ thống", 2)
for f in ["Quản lý độc giả, nhân viên, quản trị viên.",
          "Quản lý 5 loại tài liệu khác nhau.",
          "Mượn – trả tài liệu (kiểm tra điều kiện, tính hạn mượn, tính tiền phạt trễ hạn).",
          "Nhập thêm tài liệu (tính tổng chi phí) và xuất/thanh lý (tính doanh thu thanh lý).",
          "Theo dõi phiếu mượn đang mở và lịch sử giao dịch."]:
    bullet(f)

doc.add_page_break()

# ==================== CHUONG 3 ====================
heading("CHƯƠNG 3. CÀI ĐẶT CÁC KHỐI CHỨC NĂNG (CODE + KẾT QUẢ)", 1)
para("Chương này đi vào từng khối chức năng chính: trình bày đoạn code tiêu biểu và kết quả "
     "chạy thực tế trên console.")

# 3.1 Date
heading("3.1. Lớp tiện ích Date", 2)
para("Lớp Date xử lý ngày tháng: kiểm tra hợp lệ, cộng ngày, tính số ngày giữa hai mốc – làm "
     "nền tảng cho tính hạn mượn và tiền phạt.")
code_block('''class Date {
    int day, month, year;
public:
    Date(int d, int m, int y) {
        if (!isValid(d, m, y))
            throw invalid_argument("Ngay khong hop le");
        day = d; month = m; year = y;
    }
    Date addDays(int days) const;          // cong them so ngay
    int  daysUntil(const Date &other) const; // so ngay den moc khac
    bool operator<(const Date &o) const;   // so sanh ngay
};''', "Code 3.1. Lớp Date (rút gọn)")

# 3.2 Con nguoi
heading("3.2. Khối quản lý con người", 2)
para("Lớp cha trừu tượng Nguoi đóng gói thông tin cá nhân; NguoiDoc bổ sung hạn thẻ và giới "
     "hạn mượn; SinhVien / VienChuc override quy tắc hạn mượn.")
code_block('''class Nguoi {                 // LOP TRUU TUONG
protected:
    string maNguoi, hoTen, gioiTinh, soDienThoai, diaChi;
    Date   ngaySinh;
public:
    virtual string loaiNguoi() const = 0;      // ham ao thuan tuy
    virtual void   hienThiThongTin() const = 0;
};

class NguoiDoc : public Nguoi {   // ke thua + con truu tuong
    Date ngayHetHan; int soSachDangMuon, gioiHanSachMuon;
public:
    bool duocPhepMuonSach() const {
        return soSachDangMuon < gioiHanSachMuon
               && Date::today() <= ngayHetHan;
    }
    virtual int hanMuonGiaoTrinh()   const = 0;  // da hinh
    virtual int hanMuonSachThamKhao()const = 0;
    virtual int hanMuonSachKhac()    const = 0;
};

class SinhVien : public NguoiDoc {  // gioi han 5 sach
    int hanMuonSachThamKhao() const override { return 15; }
    int hanMuonSachKhac()     const override { return 7;  }
};
class VienChuc : public NguoiDoc {  // gioi han 10 sach
    int hanMuonSachThamKhao() const override { return 90; }
    int hanMuonSachKhac()     const override { return 30; }
};''', "Code 3.2. Phân cấp lớp con người")
para("Kết quả hiển thị danh sách độc giả (chức năng 6):")
output_block('''[Sinh vien]
Ma: SV001
Ho ten: Nguyen Van An
Ngay het han: 31/12/2026
So sach dang muon: 0/5
Trang thai the: Con han
MSSV: 22110001
------------------------------
[Vien chuc]
Ma: VC001
Ho ten: Tran Thi Binh
So sach dang muon: 0/10
Trang thai the: Con han
MSCB: CB2026''', "Kết quả 3.2. Danh sách độc giả (SV giới hạn 5, VC giới hạn 10)")

# 3.3 Tai lieu
heading("3.3. Khối quản lý tài liệu và đa hình tính hạn mượn", 2)
para("Lớp cha trừu tượng TaiLieu khai báo hàm ảo tinhHanMuon(). Mỗi lớp con quyết định hạn "
     "mượn bằng cách gọi ngược về quy tắc của độc giả – đây là điểm thể hiện đa hình rõ nhất.")
code_block('''class TaiLieu {               // LOP TRUU TUONG
    int soLuong; string trangThai;
public:
    bool capNhatSoLuong(int d) {         // dong goi: khong cho am
        if (soLuong + d < 0) return false;
        soLuong += d;
        trangThai = soLuong > 0 ? "Con san" : "Het";
        return true;
    }
    virtual bool coTheMuonVe() const { return true; }
    virtual int  tinhHanMuon(const NguoiDoc&) const = 0;
};

// Cung tinhHanMuon nhung moi loai goi mot quy tac khac nhau:
int GiaoTrinh::tinhHanMuon(const NguoiDoc &nd) const {
    return nd.hanMuonGiaoTrinh();      // 180 ngay
}
int SachThamKhao::tinhHanMuon(const NguoiDoc &nd) const {
    return nd.hanMuonSachThamKhao();   // SV 15 / VC 90
}
bool BaoTapChi::coTheMuonVe() const { return false; } // chi doc tai cho''',
           "Code 3.3. Đa hình trong tính hạn mượn")
para("Kết quả hiển thị tài liệu (chức năng 7) – mỗi loại hiển thị trường riêng:")
output_block('''[Giao trinh]
Ma tai lieu: GT001
Ten tai lieu: Lap trinh huong doi tuong C++
So luong: 8        Trang thai: Con san
Ma mon hoc: OOP101   Bo mon: Khoa hoc may tinh
------------------------------
[Bao tap chi]
Ma tai lieu: BC001
Ten tai lieu: Tap chi Khoa hoc Tre
So phat hanh: 6      Thang phat hanh: 6
Bao/tap chi chi doc tai cho.''', "Kết quả 3.3. Hiển thị tài liệu theo từng loại")

# 3.4 Giao dich
heading("3.4. Khối giao dịch (mượn – trả – nhập – xuất)", 2)
para("Khi mượn, hệ thống kiểm tra điều kiện, tính ngày hẹn trả theo đa hình, giảm tồn kho và "
     "ghi nhận giao dịch. Khi trả, tính tiền phạt trễ hạn (5.000đ/ngày).")
code_block('''void muonTaiLieu(...) {
    if (!nguoiDoc->duocPhepMuonSach()) { /* tu choi */ return; }
    if (!taiLieu->coTheMuonVe())       { /* chi doc tai cho */ return; }
    if (taiLieu->getSoLuong() <= 0)    { /* het sach */ return; }

    int soNgayMuon  = taiLieu->tinhHanMuon(*nguoiDoc);  // DA HINH
    Date ngayHenTra = ngayMuon.addDays(soNgayMuon);
    taiLieu->capNhatSoLuong(-1);
    nguoiDoc->tangSoSachMuon();
    dsGiaoDich.them(make_shared<Muon>(...));
}

double Tra::tinhTienPhat() const {           // phat tre han
    int soNgayTre = max(0, ngayHenTra.daysUntil(ngayGiaoDich));
    return soNgayTre * 5000.0;
}''', "Code 3.4. Nghiệp vụ mượn và tính tiền phạt")
para("Kết quả khi SinhVien SV001 mượn Giáo trình GT001 (chức năng 8):")
output_block('''Chon chuc nang: 8
Ma nguoi doc: SV001
Ma tai lieu: GT001
Ma nhan vien thuc hien: NV001
Ngay muon (bo trong = hom nay):
Muon thanh cong. Han muon: 180 ngay, hen tra: 14/12/2026.''',
             "Kết quả 3.4. Mượn giáo trình → đa hình trả về hạn 180 ngày")

# 3.5 Template
heading("3.5. Lớp khuôn mẫu DanhSach<T>", 2)
para("Một lớp template dùng chung cho mọi loại danh sách, hỗ trợ thêm (chống trùng mã), tìm, "
     "xóa và hiển thị đa hình.")
code_block('''template <typename T>
class DanhSach {
    vector<shared_ptr<T>> items;
public:
    bool them(const shared_ptr<T> &item) {
        if (!item || timTheoMa(item->getMa()) != nullptr)
            return false;                 // chong trung ma
        items.push_back(item);
        return true;
    }
    shared_ptr<T> timTheoMa(const string &ma) const;
    void hienThiTatCa() const {           // goi da hinh
        for (auto &it : items) it->hienThiThongTin();
    }
};

DanhSach<Nguoi>    dsNguoi;     // dung lai cho nhieu kieu
DanhSach<TaiLieu>  dsTaiLieu;
DanhSach<GiaoDich> dsGiaoDich;''', "Code 3.5. Lớp template DanhSach<T>")

# 3.6 Main
heading("3.6. Chương trình chính và menu điều khiển", 2)
para("Hàm main nạp dữ liệu mẫu và hiển thị menu 13 chức năng, điều phối tới các khối trên.")
output_block('''========== HE THONG QUAN LY THU VIEN ==========
1. Them nguoi doc          8. Muon tai lieu
2. Them nhan vien          9. Tra tai lieu
3. Them quan tri vien      10. Nhap them tai lieu
4. Them tai lieu           11. Xuat/thanh ly tai lieu
5. Hien thi tat ca nguoi   12. Hien thi giao dich
6. Hien thi danh sach doc gia  13. Phieu muon dang mo
7. Hien thi tai lieu       0. Thoat
Chon chuc nang:''', "Kết quả 3.6. Menu chính của chương trình")

heading("3.7. Lưu đồ thuật toán (Flowchart) các nghiệp vụ chính", 2)
para("Phần này trình bày lưu đồ thuật toán của ứng dụng web cho luồng tổng quát và hai nghiệp "
     "vụ quan trọng nhất là mượn và trả tài liệu.")

heading("a) Luồng tổng quát: đăng nhập và phân quyền", 3)
para("Người dùng đăng nhập (hoặc đăng ký độc giả mới); hệ thống xác thực tài khoản và phân "
     "quyền theo role: độc giả chỉ vào Trang độc giả + Tài liệu, còn admin/staff vào toàn bộ "
     "trang quản lý.")
add_image("fig_flow_tongquat.png",
          "Hình 3.1. Lưu đồ đăng nhập và phân quyền theo vai trò", width=5.6)

heading("b) Nghiệp vụ mượn tài liệu", 3)
para("Hệ thống kiểm tra lần lượt 5 điều kiện trước khi cho mượn; chỉ khi tất cả thỏa thì mới "
     "tính hạn mượn (đa hình theo loại tài liệu/độc giả), giảm tồn kho và ghi giao dịch.")
add_image("fig_flow_muon.png",
          "Hình 3.2. Lưu đồ nghiệp vụ mượn tài liệu (handleBorrowSubmit)", width=5.9)

heading("c) Nghiệp vụ trả tài liệu", 3)
para("Khi trả, hệ thống tính số ngày trễ và tiền phạt (5.000đ/ngày), cộng lại tồn kho, giảm "
     "số sách đang mượn của độc giả và ghi giao dịch trả.")
add_image("fig_flow_tra.png",
          "Hình 3.3. Lưu đồ nghiệp vụ trả tài liệu (handleReturnSubmit)", width=5.8)

doc.add_page_break()

# ==================== CHUONG 4 ====================
heading("CHƯƠNG 4. KẾT LUẬN", 1)
heading("4.1. Kết quả đạt được", 2)
bullet("Xây dựng hoàn chỉnh hệ thống quản lý thư viện theo mô hình OOP với 18 lớp.")
bullet("Áp dụng đầy đủ 4 đặc tính OOP cộng thêm template và smart pointer.")
bullet("Xử lý đúng các nghiệp vụ mượn – trả – nhập – xuất và kiểm tra dữ liệu đầu vào.")
bullet("Có thêm bản demo web minh họa trực quan, deploy lên GitHub Pages.")
heading("4.2. Hướng phát triển", 2)
bullet("Lưu dữ liệu xuống file/CSDL thay vì chỉ lưu trong bộ nhớ.")
bullet("Thêm thống kê, báo cáo và tìm kiếm nâng cao.")
bullet("Phân quyền chi tiết hơn và cảnh báo tự động khi sắp tới hạn trả.")

doc.save(OUT)
print("DA TAO:", OUT)
