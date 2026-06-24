# -*- coding: utf-8 -*-
"""Sinh file Word rieng cho Chuong 5 - Ket luan va huong phat trien.

Dung lai style cua build_report.py (Times New Roman 13, heading xanh, bang header xanh)
de co the dan thang vao bao cao / tieu luan.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\LTHDT_OOP\TL CK\Chuong5_KetLuan_HuongPhatTrien.docx"

doc = Document()

# ---------- Style co ban ----------
base = doc.styles["Normal"]
base.font.name = "Times New Roman"
base.font.size = Pt(13)
base._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
doc.styles["Normal"].paragraph_format.space_after = Pt(6)
doc.styles["Normal"].paragraph_format.line_spacing = 1.4


def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def para(text="", size=13, bold=False, italic=False, align=None, color=None,
         space_before=0, space_after=6, font="Times New Roman"):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def runs_para(segments, style=None, space_after=6, space_before=0):
    """Mot doan gom nhieu run (de bold tung cum)."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for text, bold in segments:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(13)
    return p


def bullet(segments):
    """segments: list[(text, bold)] hoac 1 chuoi thuong."""
    if isinstance(segments, str):
        segments = [(segments, False)]
    return runs_para(segments, style="List Bullet", space_after=3)


def heading(text, level):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # trang den
    return p


def make_table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "000000")  # nen den, chu trang -> trang den
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11.5)
    if widths:
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl


# ==================== CHUONG 5 ====================
heading("CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
para("Chương này tổng kết lại toàn bộ quá trình xây dựng đề tài Hệ thống quản lý thư viện theo "
     "mô hình lập trình hướng đối tượng (OOP): những kết quả đã đạt được, mức độ áp dụng các đặc "
     "tính OOP, những hạn chế còn tồn tại và các hướng phát triển tiếp theo. Phần đánh giá bám "
     "sát trực tiếp mã nguồn đã hiện thực (main.cpp, bản web web/app.js và web server "
     "web_server.cpp).")

# ---------- 5.1 ----------
heading("5.1. Kết luận", 2)

heading("5.1.1. Kết quả đạt được", 3)
para("Đề tài đã hoàn thành một hệ thống quản lý thư viện mô phỏng tương đối đầy đủ nghiệp vụ "
     "thực tế, tổ chức hoàn toàn theo hướng đối tượng. Cụ thể:")
bullet([("Mô hình hóa thành 3 cây kế thừa rõ ràng ", True),
        ("với tổng cộng 18 lớp trong main.cpp, xoay quanh ba nhóm thực thể của một thư viện: "
         "Con người (Nguoi → NguoiDoc → SinhVien/VienChuc; NhanVien; QuanTriVien), Tài liệu "
         "(TaiLieu → Sach → GiaoTrinh/SachThamKhao/SachKhac; BaoTapChi; BaiNghienCuu) và Giao "
         "dịch (GiaoDich → Muon/Tra/Nhap/Xuat).", False)])
bullet([("Xây dựng đầy đủ các nghiệp vụ cốt lõi ", True),
        ("qua menu 13 chức năng trong hàm main(): thêm độc giả/nhân viên/quản trị viên, thêm 5 "
         "loại tài liệu, hiển thị danh sách, mượn – trả, nhập – xuất/thanh lý, xem giao dịch và "
         "phiếu mượn đang mở.", False)])
bullet([("Xử lý đúng các ràng buộc nghiệp vụ quan trọng: ", True),
        ("kiểm tra điều kiện mượn trong duocPhepMuonSach() (còn hạn thẻ và chưa vượt giới hạn — "
         "Sinh viên tối đa 5 cuốn, Viên chức tối đa 10 cuốn); chặn mượn về tài liệu chỉ đọc tại "
         "chỗ qua coTheMuonVe(); cập nhật tồn kho an toàn trong capNhatSoLuong() (không cho âm, "
         "tự đổi trạng thái Còn sẵn/Hết); tính tiền phạt trễ hạn trong tinhTienPhat() theo công "
         "thức 5.000đ/ngày.", False)])
bullet([("Kiểm soát dữ liệu đầu vào ", True),
        ("bằng các hàm tiện ích readInt(), readDouble(), readDate(), readNonEmpty() và lớp Date "
         "(kiểm tra ngày hợp lệ, năm nhuận, định dạng dd/mm/yyyy), tránh nhập sai gây lỗi.", False)])
bullet([("Phát triển thêm bản demo web trực quan ", True),
        ("trong thư mục web/ chạy bằng localStorage, có phân quyền 3 vai trò (admin/staff/reader) "
         "và đã được triển khai tự động lên GitHub Pages.", False)])

heading("5.1.2. Đánh giá việc áp dụng các đặc tính OOP", 3)
para("Đề tài không chỉ “có dùng” mà còn thể hiện rõ vai trò của từng đặc tính OOP. Bảng sau đối "
     "chiếu mỗi đặc tính với vị trí áp dụng cụ thể trong mã nguồn:")
make_table(
    ["Đặc tính OOP", "Thể hiện trong mã nguồn"],
    [
        ["Trừu tượng\n(Abstraction)",
         "Ba lớp cha Nguoi, TaiLieu, GiaoDich đều có hàm ảo thuần túy (= 0) như "
         "hienThiThongTin(), tinhHanMuon(), thucHienGiaoDich() nên không thể khởi tạo trực "
         "tiếp, chỉ làm khung thiết kế."],
        ["Đóng gói\n(Encapsulation)",
         "Thuộc tính để private/protected; truy xuất qua getter/setter. Ràng buộc được gói "
         "trong phương thức — capNhatSoLuong() chặn số lượng âm, duocPhepMuonSach() gói toàn "
         "bộ logic kiểm tra điều kiện mượn."],
        ["Kế thừa\n(Inheritance)",
         "Kế thừa phân cấp nhiều cấp: Nguoi → NguoiDoc → SinhVien/VienChuc và TaiLieu → Sach "
         "→ GiaoTrinh/... Lớp con tái sử dụng toàn bộ lớp cha và chỉ thêm phần riêng (SinhVien "
         "thêm mssv, VienChuc thêm mscb)."],
        ["Đa hình\n(Polymorphism)",
         "Rõ nhất ở tinhHanMuon(): cùng lời gọi taiLieu->tinhHanMuon(*nguoiDoc) nhưng cho hạn "
         "khác nhau theo loại tài liệu và loại độc giả (Giáo trình 180; Sách tham khảo SV 15/VC "
         "90; Sách khác SV 7/VC 30). Mỗi lớp con cũng override hienThiThongTin()."],
        ["Khuôn mẫu\n(Template)",
         "Lớp DanhSach<T> viết một lần, dùng lại cho DanhSach<Nguoi>, DanhSach<NguoiDoc>, "
         "DanhSach<TaiLieu>, DanhSach<GiaoDich>; có sẵn them() (chống trùng mã), timTheoMa(), "
         "xoaTheoMa(), hienThiTatCa()."],
        ["Quản lý bộ nhớ an toàn",
         "Dùng shared_ptr<T> cho mọi đối tượng và khai báo destructor ảo (virtual ~Nguoi(), "
         "virtual ~TaiLieu(), virtual ~GiaoDich()) để hủy đúng đối tượng lớp con, tránh rò rỉ."],
    ],
    widths=[1.7, 4.8],
)
para("Nhận xét chung: Hệ thống đã hiện thực đầy đủ 4 trụ cột của OOP (trừu tượng, đóng gói, kế "
     "thừa, đa hình) cộng thêm template và smart pointer. Cách tổ chức lớp giúp mã nguồn dễ đọc, "
     "dễ mở rộng: thêm một loại tài liệu mới chỉ cần tạo một lớp con của TaiLieu và override "
     "tinhHanMuon()/coTheMuonVe() mà không sửa code cũ — đúng tinh thần “mở để mở rộng, đóng để "
     "sửa đổi”.", italic=True)

heading("5.1.3. Hạn chế còn tồn tại", 3)
para("Bên cạnh kết quả đạt được, qua rà soát mã nguồn nhóm nhận thấy hệ thống còn một số hạn chế:")
bullet([("Chưa lưu trữ dữ liệu lâu dài (bản console): ", True),
        ("toàn bộ dữ liệu của main.cpp nằm trong các DanhSach<T> và vector trên RAM; khi thoát "
         "chương trình mọi thay đổi đều mất, lần chạy sau phải nạp lại từ napDuLieuMau().", False)])
bullet([("Bản web chỉ lưu cục bộ: ", True),
        ("dữ liệu lưu trong localStorage của trình duyệt, không đồng bộ giữa các máy và sẽ mất "
         "nếu xóa dữ liệu duyệt web.", False)])
bullet([("Bản console và bản web tách rời: ", True),
        ("dùng chung mô hình thiết kế nhưng là hai codebase độc lập, không chia sẻ dữ liệu.", False)])
bullet([("Bảo mật còn đơn giản: ", True),
        ("mật khẩu (ví dụ admin/admin123 trong QuanTriVien) lưu ở dạng văn bản thuần, chưa "
         "mã hóa/băm; phần ******** chỉ là che hiển thị.", False)])
bullet([("Một số tham số bị cố định cứng: ", True),
        ("mức phạt 5.000đ/ngày và các mốc hạn mượn viết thẳng trong code, chưa cấu hình được.", False)])
bullet([("Thiếu thống kê – báo cáo: ", True),
        ("mới dừng ở liệt kê danh sách, chưa có thống kê sách mượn nhiều, độc giả trễ hạn, "
         "doanh thu nhập/xuất...", False)])
bullet([("Phụ thuộc nền tảng (web server C++): ", True),
        ("web_server.cpp dùng Winsock (ws2_32) nên chỉ chạy được trên Windows.", False)])
bullet([("Chưa có kiểm thử tự động: ", True),
        ("mới dừng ở kiểm tra cú pháp (node --check) và chạy thử thủ công, chưa có unit test cho "
         "logic nghiệp vụ.", False)])

# ---------- 5.2 ----------
heading("5.2. Hướng phát triển", 2)
para("Từ những hạn chế trên, nhóm đề xuất các hướng phát triển tiếp theo, sắp xếp theo mức độ "
     "ưu tiên:")

heading("5.2.1. Lưu trữ và quản lý dữ liệu", 3)
bullet("Lưu dữ liệu xuống file hoặc cơ sở dữ liệu thay vì chỉ giữ trên RAM: xuất/nhập file "
       "CSV/JSON cho bản console, hoặc tích hợp SQLite/MySQL để dữ liệu tồn tại qua các lần chạy.")
bullet("Thêm chức năng sao lưu – phục hồi dữ liệu để tránh mất mát.")

heading("5.2.2. Hợp nhất kiến trúc console – web", 3)
bullet("Đưa lõi C++ thành backend thật sự: nâng cấp web_server.cpp thành API (REST/JSON) để bản "
       "web gọi trực tiếp, dùng chung một nguồn dữ liệu thay vì localStorage cục bộ.")
bullet("Thay Winsock bằng thư viện socket đa nền tảng (hoặc framework như Crow/Drogon) để web "
       "server chạy được trên cả Linux/macOS.")

heading("5.2.3. Bổ sung nghiệp vụ và tiện ích", 3)
bullet("Thống kê – báo cáo: sách được mượn nhiều nhất, danh sách độc giả đang trễ hạn, tổng chi "
       "phí nhập và doanh thu thanh lý theo kỳ.")
bullet("Cảnh báo tự động khi sách sắp đến hạn hoặc đã quá hạn trả (thông báo/email).")
bullet("Đặt trước (reservation): cho phép độc giả đặt giữ chỗ khi sách đang được mượn hết.")
bullet("Quản lý thanh toán tiền phạt và cho cấu hình linh hoạt mức phạt, hạn mượn thay vì "
       "hard-code.")
bullet("Mở rộng tìm kiếm – lọc nâng cao (đã có ở bản web theo loại/chủ đề) sang bản console.")

heading("5.2.4. Bảo mật và chất lượng phần mềm", 3)
bullet("Mã hóa/băm mật khẩu (ví dụ bcrypt/SHA-256 + salt) và quản lý phiên đăng nhập an toàn.")
bullet("Phân quyền chi tiết hơn ở mức từng chức năng thay vì chỉ theo vai trò.")
bullet("Viết unit test cho các phần logic quan trọng (Date, tinhHanMuon(), tinhTienPhat(), "
       "capNhatSoLuong()) để bảo đảm tính đúng đắn khi mở rộng.")

# ---------- 5.3 ----------
heading("5.3. Tổng kết", 2)
para("Qua đề tài, nhóm đã vận dụng được trọn vẹn kiến thức môn Lập trình hướng đối tượng vào "
     "một bài toán thực tế: phân tích – thiết kế hệ thống thành các lớp, áp dụng đầy đủ trừu "
     "tượng, đóng gói, kế thừa, đa hình, kết hợp template và smart pointer để tạo ra một chương "
     "trình rõ ràng, an toàn bộ nhớ và dễ mở rộng. Sản phẩm gồm cả bản console C++ và bản web "
     "demo đã chạy trực tuyến, đáp ứng tốt yêu cầu của một đồ án cuối kỳ. Những hạn chế còn lại "
     "(lưu trữ, bảo mật, thống kê) chính là tiền đề cho các hướng phát triển tiếp theo nếu nâng "
     "cấp hệ thống lên mức ứng dụng hoàn chỉnh.")

doc.save(OUT)
print("DA TAO:", OUT)
