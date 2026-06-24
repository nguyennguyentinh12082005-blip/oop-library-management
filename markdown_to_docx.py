"""
Chuyen doi file Bao_Cao_OOP.md sang file Word (.docx) voi format chuyen nghiep.
KHONG tao trang bia va muc luc.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Paths -----------------------------------------------------------------
MD_FILE = r"c:\Users\tcmti\Downloads\Bao_Cao_OOP.md"
OUTPUT_DOCX = r"c:\Users\tcmti\Downloads\Bao_Cao_OOP_v2.docx"
ASSETS_DIR = r"e:\LTHDT_OOP\TL CK\report_assets"


# --- Helpers ---------------------------------------------------------------

def read_md(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        rPr = run.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)
    return run


def apply_inline_formatting(paragraph, text, default_size=13, default_font="Times New Roman"):
    """Xu ly bold (**), italic (*), va inline code (`) trong text."""
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):
            add_formatted_text(paragraph, match.group(2), bold=True, size=default_size, font_name=default_font)
        elif match.group(3):
            add_formatted_text(paragraph, match.group(3), italic=True, size=default_size, font_name=default_font)
        elif match.group(4):
            run = add_formatted_text(paragraph, match.group(4), size=10, font_name="Consolas")
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        else:
            add_formatted_text(paragraph, match.group(5), size=default_size, font_name=default_font)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


# --- Parse markdown --------------------------------------------------------

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r'^\|[\s:\-]+\|', line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    return rows


def clean_text(text):
    text = text.replace("\\-", "-").replace("\\<", "<").replace("\\>", ">")
    text = text.replace("\\=", "=").replace("\\&", "&")
    text = re.sub(r'\\([^\\])', r'\1', text)
    return text.strip()


def parse_md_blocks(md_text):
    blocks = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(('code_block', lang, "\n".join(code_lines)))
            i += 1
            continue

        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = re.sub(r'^\*\*(.+?)\*\*$', r'\1', text)
            text = clean_text(text)
            blocks.append(('heading', level, text))
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                blocks.append(('table', rows))
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\\?\.\s+(.+)$', stripped)
        if num_match:
            text = clean_text(num_match.group(2))
            blocks.append(('numbered_item', num_match.group(1), text))
            i += 1
            continue

        # Bullet (unicode or markdown)
        if stripped.startswith(chr(0x2022)) or stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped.lstrip(chr(0x2022)).lstrip("-").lstrip("*").strip()
            text = clean_text(text)
            blocks.append(('bullet_item', text))
            i += 1
            continue

        # Tree-like lines
        if any(c in stripped for c in ['\u251c', '\u2514', '\u2502']):
            blocks.append(('tree_line', stripped))
            i += 1
            continue

        # Bold-only line
        bold_match = re.match(r'^\*\*(.+?)\*\*$', stripped)
        if bold_match:
            text = clean_text(bold_match.group(1))
            blocks.append(('bold_paragraph', text))
            i += 1
            continue

        # Regular paragraph
        para_text = stripped
        while i + 1 < len(lines) and lines[i].rstrip().endswith("  "):
            i += 1
            next_line = lines[i].strip()
            if next_line:
                para_text += "\n" + next_line
        para_text = clean_text(para_text)
        blocks.append(('paragraph', para_text))
        i += 1

    return blocks


# --- Build Word Document ---------------------------------------------------

def create_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for level in range(1, 5):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_font = h_style.font
            h_font.name = 'Times New Roman'
            h_font.bold = True
            h_font.color.rgb = RGBColor(0, 0, 0)
            h_rPr = h_style.element.get_or_add_rPr()
            h_rFonts = h_rPr.find(qn('w:rFonts'))
            if h_rFonts is None:
                h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/>')
                h_rPr.append(h_rFonts)
            else:
                h_rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            if level == 1:
                h_font.size = Pt(16)
            elif level == 2:
                h_font.size = Pt(14)
            elif level == 3:
                h_font.size = Pt(13)

    return doc


def add_table_to_doc(doc, rows):
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]

            ct = cell_text.replace('\x0b', '\n').replace('<br>', '\n').strip()
            ct = re.sub(r'\*\*(.+?)\*\*', r'\1', ct)
            ct = ct.replace('\\', '')

            if row_idx == 0:
                set_cell_shading(cell, "1F4E79")
                add_formatted_text(p, ct, bold=True, size=11, color=(255, 255, 255), font_name="Times New Roman")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if row_idx % 2 == 0:
                    set_cell_shading(cell, "D6E4F0")

                bold_match = re.match(r'^(.+?\(L.+?\))\s*(.*)', ct, re.DOTALL)
                if bold_match and col_idx == 0:
                    add_formatted_text(p, bold_match.group(1), bold=True, size=10, font_name="Times New Roman")
                    if bold_match.group(2):
                        add_formatted_text(p, " " + bold_match.group(2), size=10, font_name="Times New Roman")
                else:
                    add_formatted_text(p, ct, size=10, font_name="Times New Roman")

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            tcPr.append(vAlign)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tblPr.append(tblW)

    doc.add_paragraph()


def add_image_if_exists(doc, filename, caption=None, width=None):
    filepath = os.path.join(ASSETS_DIR, filename)
    if os.path.exists(filepath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=width or Inches(5.5))

        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(cap_p, caption, italic=True, size=11, font_name="Times New Roman", color=(100, 100, 100))
            set_paragraph_spacing(cap_p, before=4, after=12)
        return True
    return False


# --- Image mapping for sections --------------------------------------------

IMAGE_MAP = {
    "3.1": [
        ("fig_flowchart_system.png", "Hinh 3.1: So do khoi tong quan he thong"),
    ],
    "3.2": [
        ("fig_dfd_level1.png", "Hinh 3.2: So do chi tiet cac khoi chuc nang"),
        ("fig_flow_muon.png", "Hinh 3.3: Flowchart quy trinh muon tai lieu"),
        ("fig_flow_tra.png", "Hinh 3.4: Flowchart quy trinh tra tai lieu"),
    ],
    "2.1": [
        ("fig_object_model.png", "Hinh 2.1: So do mo hinh doi tuong he thong"),
    ],
}

# Keep a copy so we can consume entries
_image_map_copy = {}


def build_docx(md_text, output_path):
    doc = create_document()

    # NO cover page, NO table of contents - start directly with content

    blocks = parse_md_blocks(md_text)

    # Make a consumable copy of image map
    img_map = {}
    for k, v in IMAGE_MAP.items():
        img_map[k] = list(v)

    tree_block_lines = []
    current_section = ""

    for idx, block in enumerate(blocks):
        btype = block[0]

        # Flush tree block
        if btype != 'tree_line' and tree_block_lines:
            p = doc.add_paragraph()
            tree_text = "\n".join(tree_block_lines)
            add_formatted_text(p, tree_text, size=11, font_name="Consolas")
            set_paragraph_spacing(p, before=2, after=6, line_spacing=1.0)
            tree_block_lines = []

        if btype == 'heading':
            level, text = block[1], block[2]

            # Track section number
            sec_match = re.match(r'^.*?(\d+\.\d+)', text)
            if sec_match:
                current_section = sec_match.group(1)

            heading_p = doc.add_heading(text, level=min(level, 4))
            heading_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading_p.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
                rPr = run.element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/>')
                    rPr.append(rFonts)
                else:
                    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            set_paragraph_spacing(heading_p, before=12, after=6)

            # Insert images after relevant section heading
            if current_section in img_map:
                for img_file, caption in img_map[current_section]:
                    add_image_if_exists(doc, img_file, caption)
                del img_map[current_section]

        elif btype == 'table':
            rows = block[1]
            add_table_to_doc(doc, rows)

        elif btype == 'paragraph':
            text = block[1]
            p = doc.add_paragraph()
            apply_inline_formatting(p, text, default_size=13)
            set_paragraph_spacing(p, before=4, after=4)

        elif btype == 'bold_paragraph':
            text = block[1]
            p = doc.add_paragraph()
            add_formatted_text(p, text, bold=True, size=13, font_name="Times New Roman")
            set_paragraph_spacing(p, before=6, after=4)

        elif btype == 'bullet_item':
            text = block[1]
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text, default_size=13)
            set_paragraph_spacing(p, before=2, after=2)

        elif btype == 'numbered_item':
            num, text = block[1], block[2]
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text, default_size=13)
            set_paragraph_spacing(p, before=2, after=2)

        elif btype == 'tree_line':
            tree_block_lines.append(block[1])

        elif btype == 'code_block':
            lang, code = block[1], block[2]
            p = doc.add_paragraph()
            add_formatted_text(p, code, size=10, font_name="Consolas")
            set_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                f'<w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                f'<w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            pPr.append(shd)

    # Flush remaining tree lines
    if tree_block_lines:
        p = doc.add_paragraph()
        tree_text = "\n".join(tree_block_lines)
        add_formatted_text(p, tree_text, size=11, font_name="Consolas")
        set_paragraph_spacing(p, before=2, after=6, line_spacing=1.0)

    doc.save(output_path)
    print(f"[OK] Da xuat thanh cong: {output_path}")
    print(f"   Kich thuoc: {os.path.getsize(output_path) / 1024:.1f} KB")


# --- Main ------------------------------------------------------------------

if __name__ == "__main__":
    if not os.path.exists(MD_FILE):
        print(f"[ERROR] Khong tim thay file: {MD_FILE}")
        exit(1)

    md_content = read_md(MD_FILE)
    print(f"[READ] Doc file MD: {len(md_content)} ky tu")
    print(f"[DIR] Thu muc anh: {ASSETS_DIR}")
    if os.path.exists(ASSETS_DIR):
        imgs = [f for f in os.listdir(ASSETS_DIR) if f.endswith('.png')]
        print(f"   Cac anh co san: {', '.join(imgs)}")
    else:
        print("   Khong tim thay thu muc anh")
    print("[WRITE] Dang tao file Word...")

    build_docx(md_content, OUTPUT_DOCX)
